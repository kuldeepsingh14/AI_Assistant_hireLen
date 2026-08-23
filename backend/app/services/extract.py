"""Turn an uploaded resume file into clean plain text."""
from __future__ import annotations

import io
import re

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}


class UnsupportedFile(ValueError):
    pass


def extract_text(filename: str, raw: bytes) -> str:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == ".pdf":
        text = _from_pdf(raw)
    elif ext == ".docx":
        text = _from_docx(raw)
    elif ext in (".txt", ".md"):
        text = raw.decode("utf-8", errors="replace")
    else:
        raise UnsupportedFile(
            f"Unsupported file type '{ext or filename}'. Use one of: {', '.join(sorted(SUPPORTED))}"
        )
    return _clean(text)


def _from_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # a single broken page shouldn't kill the upload
            pages.append("")
    return "\n".join(pages)


def _from_docx(raw: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ").replace("\u2022", "- ").replace("\ufb01", "fi")
    # PDF extractors love to emit "l i k e   t h i s" or stray ligatures; collapse runs of spaces.
    text = re.sub(r"[ \t]+", " ", text)
    # Drop lines that are pure page furniture ("Page 1 of 2").
    lines = [
        ln.strip()
        for ln in text.split("\n")
        if not re.fullmatch(r"(page\s*)?\d+(\s*(of|/)\s*\d+)?", ln.strip(), re.I)
    ]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_filename(filename: str, fallback: str = "resume") -> str:
    """Reduce an uploaded filename to a harmless basename.

    An upload's filename is attacker-controlled, and both storing and later
    serving the file join it onto a directory - so "../../app/main.py" would
    escape the data folder entirely. Strip every path separator and keep only
    characters that cannot mean anything to a filesystem.
    """
    # Handle both separators regardless of host OS, plus Windows drive prefixes.
    base = re.split(r"[\\/]", filename.strip())[-1]
    base = re.sub(r"^[A-Za-z]:", "", base)

    stem, _, ext = base.rpartition(".")
    if not stem:  # no dot at all
        stem, ext = base, ""

    stem = re.sub(r"[^A-Za-z0-9._ -]+", "_", stem).strip(" ._-")
    ext = re.sub(r"[^A-Za-z0-9]+", "", ext).lower()

    if not stem:
        stem = fallback
    stem = stem[:80]

    return f"{stem}.{ext}" if ext else stem
